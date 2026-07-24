import io
from typing import List, Dict, Any
from fpdf import FPDF
from docx import Document as DocxDocument

class ExporterService:
    @staticmethod
    def generate_pdf_notes(project_name: str, notes: List[Dict[str, Any]]) -> bytes:
        """
        Creates a clean PDF document containing all generated notes sorted by outline modules/topics.
        """
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Cover Page
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(0, 40, "", ln=True) # spacer
        pdf.cell(0, 15, txt="StudyForge AI - Study Guide", ln=True, align="C")
        pdf.set_font("Helvetica", "I", 14)
        pdf.cell(0, 10, txt=f"Subject Workspace: {project_name}", ln=True, align="C")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 20, txt="Consolidated reference textbook-grounded notes", ln=True, align="C")
        pdf.add_page()

        for note in notes:
            module_title = note.get("module_title", "Module")
            topic_name = note.get("topic_name", "Topic")
            content = note.get("content", "")

            # Heading 1: Module
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(16, 185, 129) # Emerald Green color
            pdf.cell(0, 10, txt=f"Module: {module_title}", ln=True)
            pdf.set_text_color(0, 0, 0) # Reset color to black
            
            # Heading 2: Topic
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, txt=f"Topic: {topic_name}", ln=True)
            pdf.ln(2)

            # Body text parsing paragraph-by-paragraph
            pdf.set_font("Helvetica", "", 10)
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                # Clean simple markdown tags (remove bold # tags for PDF cells)
                clean_para = para.replace("**", "").replace("###", "").replace("##", "").replace("#", "").strip()
                # Remove citation tags for printable layout cleanliness
                clean_para = re_clean = "".join(c for c in clean_para if ord(c) < 128) # ASCII clean only
                pdf.multi_cell(0, 5, txt=clean_para)
                pdf.ln(3)
                
            pdf.ln(10)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)

        # Output raw bytes
        return bytes(pdf.output())

    @staticmethod
    def generate_docx_notes(project_name: str, notes: List[Dict[str, Any]]) -> bytes:
        """
        Creates a structured Word Document (.docx) containing the consolidated study guide.
        """
        doc = DocxDocument()
        
        # Cover details
        doc.add_heading(f"StudyForge AI - Study Guide", level=0)
        doc.add_paragraph(f"Subject Workspace: {project_name}")
        doc.add_paragraph("Consolidated Reference Textbook-Grounded Notes").italic = True
        doc.add_page_break()

        for note in notes:
            module_title = note.get("module_title", "Module")
            topic_name = note.get("topic_name", "Topic")
            content = note.get("content", "")

            doc.add_heading(f"Module: {module_title}", level=1)
            doc.add_heading(f"Topic: {topic_name}", level=2)
            
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                # We write standard paragraphs
                doc.add_paragraph(para)
            
            doc.add_paragraph().paragraph_format.space_after = 20

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
